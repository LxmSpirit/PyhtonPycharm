import xlrd
import os
import re
import jieba
import sys
from docx import Document
from docx.shared import Cm

URL = 'D:/Work/Python/PycharmProjects/untitled/CIMS/ALS/321.xls'
URLfolder = 'D:/Work/Python/PycharmProjects/untitled/CIMS/ALS'
#D:\Work\Python\PycharmProjects\untitled\CIMS\ALS
formaddress = 'D:/Work/Python/PycharmProjects/untitled/Rave-ALS/formdata6.txt'
fieldaddress = 'D:/Work/Python/PycharmProjects/untitled/Rave-ALS/fieldata914.txt'
URLcrf = 'D:/Work/Python/PycharmProjects/untitled/CIMS/CRFdraft/KMHH-PIBC-101_UniqueCRF_V1.0_20230116.docx'
URLcrfsave = 'D:/Work/Python/PycharmProjects/untitled/CIMS/CRFdraft/CRF1-2.docx'
URLcrfsave2 = 'D:/Work/Python/PycharmProjects/untitled/CIMS/CRFdraft/CRF1-918.docx'
URLlist = []
def readfolder():
    # 遍历指定文件夹下的所有Excel文件并执行分析操作
    for file_name in os.listdir(URLfolder):
        if file_name.endswith('.xls'):
            # 构造输入和输出文件的路径
            #input_path = os.path.join(URLfolder, file_name)
            input_path = URLfolder+"/"+file_name
            URLlist.append(input_path)

    return URLlist

def newreadform(urllist):
    dicform = {}
    for url in urllist:
        workbook = xlrd.open_workbook(url)

        sheet1 = workbook.sheet_by_name('eCRF表单')
        nrows = sheet1.nrows  # 行数
        ncols = sheet1.ncols  # 列数

        for i in range(1, nrows):
            cellOID = sheet1.cell(i, 0).value
            cellName = sheet1.cell(i, 1).value

            newOID = re.sub(r'\d+', '', cellOID)
            pattern = re.compile('[!@#$%^&*\(\)（）_+[\]{};:,，./<>?\|`~-]')
            cellName2 = re.sub(pattern, '', cellName)
            newName = jieba.lcut_for_search(cellName2)
            s = ""
            for name in newName:
                s = s + " " + name

            if newOID in dicform:
                oldv = dicform[newOID]
                dicform[newOID] = oldv + " " + s
            else:
                dicform[newOID] = s

        if "FT" in dicform:
            dicform.pop("FT")
        if "CPH" in dicform:
            dicform.pop("CPH")

    return dicform

def newreadfield(urllist):
    dicfield = {}
    for url in urllist:
        workbook = xlrd.open_workbook(url)
        # print(workbook.sheet_names())
        sheet1 = workbook.sheet_by_name('变量列表')
        nrows = sheet1.nrows  # 行数
        ncols = sheet1.ncols  # 列数
        i = 17
        while i < nrows:
            line = sheet1.cell(i, 0).value
            # print(line)
            if "--" in line:

                formlist = line.split("--")
                formoid2 = formlist[0]
                formoid = re.sub(r'\d+', '', formoid2)

                i = i + 1
                linfieid = sheet1.cell(i, 0).value
                linfiena = sheet1.cell(i, 1).value

                while len(linfieid) > 0:
                    i = i + 1
                    if i >= nrows:
                        break
                    linfieid = sheet1.cell(i, 0).value
                    linfiena = sheet1.cell(i, 1).value
                    if len(linfieid) > 0:
                        pattern = re.compile('[!？@#$%^&*\(\)（）_+[\]{};:,，./<>?\|`~-]')
                        fieldName2 = re.sub(pattern, '', linfiena)  # 处理后的fieldname
                        if (formoid in linfieid):
                            newfieldOID = linfieid.replace(formoid, "")
                        else:
                            newfieldOID = linfieid  # 处理后的fieldoid

                        newName = jieba.lcut_for_search(fieldName2)

                        s = ""
                        for name in newName:
                            s = s + " " + name

                        if newfieldOID in dicfield:
                            oldv = dicfield[newfieldOID]
                            dicfield[newfieldOID] = oldv + " " + s
                        else:
                            dicfield[newfieldOID] = s
            i = i + 1
    return dicfield


def write(data,address):
    file = open(address,mode='w',encoding='utf-8')
    for k,v in data.items():
        file.write(k+'_'+v+'\n')

def coutdic(data):
    keysum = len(data)
    dicv = {}
    onlyv={}
    dick = {}
    maximum = {}
    Prior = {}
    pv = {}
    for k in data:
        #strv = str(v)
        v = data[k]
        vlist = v.split(" ")
        while "" in vlist:
            vlist.remove("")
        for i in vlist:
            str = k+'-'+i
            if str in dicv:
                num = dicv[str]
                dicv[str] = num +1
            else:
                dicv[str] =  1
        if k in dick:
            numk = dick[k]
            dick[k] = numk+1
        else:
            dick[k] =len(vlist)
    #用于计算v
    for k in data:
        #strv = str(v)
        v = data[k]
        vlist = v.split(" ")
        while "" in vlist:
            vlist.remove("")
        for i in vlist:
            if i in onlyv:
                num = onlyv[i]
                onlyv[i] = num +1
            else:
                onlyv[i] =  1
    sumv = 0

    for a in onlyv:
        sumv = sumv+onlyv[a]

    for a in onlyv:
        num = onlyv[a] / sumv
        pv[a] = num

    for key in dicv:
        n1 = dicv[key]
        k = key.split("-")
        n2 = dick[k[0]]
        res = n1 / n2
        maximum[key] = res

    sum = 0
    for key in dick:
        n1 = dick[key]
        sum = sum +n1

    for key in dick:
        n1 = dick[key]
        res = n1 / sum
        Prior[key] = res

    return maximum,Prior,pv


def predict2(pxy,py,px,str):
    pattern = re.compile('[!@#$%^&*\(\)（）_+[\]{};:,./<>?\|`~-]')
    newstr1 = re.sub(pattern, '', str)
    newstr = jieba.lcut_for_search(newstr1)
    print(newstr)
    res = {}

    for prestr in newstr:
        for k in pxy:#训练数据集
            k1 = k.split("-")
            key = k1[0]#结果的Key
            a1 = 0
            a2 = 0
            a3 = 1
            if k1[1]==prestr:
                a1 = pxy[k]
                if k1[1] in px:
                    a3 = px[k1[1]]
                if key in py:
                    a2 = py[key]
                resu = (a1 * a2) / a3
                if key in res:  # 计算结果
                    a = res[key]
                    v = a + resu
                    res[key] = v
                else:
                    res[key] = resu

    return res

def endres(res):
    maxres = -1
    same = {}
    for k in res:
        if res[k]>maxres:
            maxres = res[k]
            resk = k
    if maxres !=-1:
        for k in res:
            if res[k] == maxres:
                same[k] = maxres
        return resk,same
    else:
        print("QAQ")
        return None

def removenull(data):
    for k in data:
        v = data[k]
        vlist = v.split(" ")
        print(vlist)
        while "" in vlist:
            vlist.remove("")
        print(vlist)
        print("---")

def readcrf():
    doc = Document(URLcrf)
    list = []
    # 读取文档中的内容
    for paragraph in doc.paragraphs:
        if paragraph.style.name == 'Heading 1' or paragraph.style.name == 'Heading 2':
            list.append(paragraph.text)#获得标题列表

    tables = doc.tables
    a = 0
    for table in tables:
        table.add_row()
        rownum = len(table.rows)
        print(rownum)
        subnum =1
        rows = rownum - 1
        if a < len(list) and len(table.rows[rows].cells)>0:
            table.rows[rownum - 1].cells[0].text = list[a]
            print(list[a])
        while len(table.rows[rows].cells)==0:
            subnum = subnum + 1
            rows  = rownum -subnum
        table.rows[rows].cells[0].text = list[a]
        a = a + 1


    doc.save(URLcrfsave)

def writecrf(pxy,py,px):
    doc = Document(URLcrfsave)
    listpar = []
    # 读取文档中的内容
    for paragraph in doc.paragraphs:
        if paragraph.style.name == 'Heading 1' or paragraph.style.name == 'Heading 2':
            listpar.append(paragraph.text)#获得标题列表
    listnew = []
    for par in listpar:
        res = predict2(pxy, py, px, par)
        a, list = endres(res)
        if len(list)>1 :
            print("多结果")
            print(list)
            sumitem=""
            for item in list:
                sumitem = sumitem +"/"+item
            newpar = par+"_"+sumitem
            listnew.append(newpar)
        else:
            print("单结果")
            print(a)
            newpar = par+"_"+a
            listnew.append(newpar)
    print("----")
    print(listnew)
    print(len(listnew))
    print(len(listpar))
    index = 0
    if len(listnew) == len(listpar):
        for paragraph in doc.paragraphs:
            if paragraph.style.name == 'Heading 1' or paragraph.style.name == 'Heading 2':
                paragraph.text =  listnew[index] # 获得标题列表
                index = index +1

    doc.save(URLcrfsave2)


if __name__ == '__main__':
    readcrf()

    urllist = readfolder()

    dicform = newreadform(urllist)
    pxy,py,px = coutdic(dicform)
    writecrf(pxy, py, px)
    #str = "PK采血"
    #res = predict2(pxy,py,px, str)
    #a,list=endres(res)
    #if len(list)>1 :
        #print("多结果")
        #print(list)
    #else:
        #print("单结果")
        #print(a)


    #dicfield = newreadfield(urllist)
    #pxy1, py1, px1 = coutdic(dicfield)
    #str = "研究终止"
    #res = predict2(pxy1, py1, px1, str)
    #a,list=endres(res)
    #print("testfield: " + str)
    #if len(list)>1 :
        #print("多结果")
        #print(list)
    #else:
        #print("单结果")
        #print(a)





