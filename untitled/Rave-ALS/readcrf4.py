from docx import Document

doc = Document('D:/Work/Python/PycharmProjects/untitled/Rave-ALS/CRF1.docx')

list = []

# 读取文档中的内容
for paragraph in doc.paragraphs:
    if paragraph.style.name == 'Heading 1' or paragraph.style.name == 'Heading 2':
        list.append(paragraph.text)


tables = doc.tables
num = len(tables)
b = len(list)

a = 0
for table in tables:
    table.add_row()
    rownum = len(table.rows)
    if a <b :
        table.rows[rownum-1].cells[0].text = list[a]

    a = a+1

doc.save('D:/Work/Python/PycharmProjects/untitled/Rave-ALS/CRF1-2.docx')