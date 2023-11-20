
from translate import Translator
import docx
from openpyxl import Workbook
#import synonyms
#synlst = synonyms.display('序号')

path = input('请输入文件路径(.docx):\n')

def read_text(path,CHN,ENC):
    # file = open(path,encoding="utf-8")
    wb = Workbook()  # 建一个新的工作簿
    sht = wb.active
    text =docx.Document(path)
    tables = text.tables
    #table = text[0]
    for i in (0,len(tables)-1):
        for row in tables[i].rows:
            for cell in row.cells:
                cell.text = cell.text.replace(CHN,ENC)
                #print(cell.text)
    for table in text.tables:
            for i, row in enumerate(table.rows):
                excel_row = []
                excel_space = []
                for cell in row.cells:
                    excel_row.append(cell.text)
                    excel_space.append(" ")
                sht.append(excel_row)
            sht.append(excel_space)
    wb.save(r'C:\Users\Lenovo\Desktop\2.xlsx')
    text.save(r'C:\Users\Lenovo\Desktop\2.docx')

CHN = input('请输入中文词汇：\n')
ENC = input('请输入英文词汇：\n')
text = read_text(path,CHN,ENC)
